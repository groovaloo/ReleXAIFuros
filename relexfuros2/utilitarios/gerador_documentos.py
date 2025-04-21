import os
from datetime import datetime
from docx import Document

def guardar_carta_formatada(texto_txt, nome="Cooper", destino="Agência Portuguesa do Ambiente", local="Lisboa"):
    pasta = "relexfuros2/registos"
    os.makedirs(pasta, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    base_nome = f"carta_{timestamp}"
    caminho_txt = os.path.join(pasta, base_nome + ".txt")
    caminho_doc = os.path.join(pasta, base_nome + ".docx")

    # Guardar .txt
    with open(caminho_txt, "w", encoding="utf-8") as f:
        f.write(texto_txt)

    # Guardar .docx
    doc = Document()
    doc.add_paragraph(f"{local}, {datetime.now().strftime('%d/%m/%Y')}\n")
    doc.add_paragraph(f"De: {nome}\n")
    doc.add_paragraph(f"Para: {destino}\n")
    doc.add_paragraph("Assunto: Denúncia relativa à execução de furo não licenciado\n")
    doc.add_paragraph(texto_txt)
    doc.add_paragraph("\nCom os melhores cumprimentos,\n")
    doc.add_paragraph(f"{nome}")
    doc.save(caminho_doc)

    return f"🗂️ Carta guardada em:\n- TXT: {caminho_txt}\n- DOCX: {caminho_doc}"
